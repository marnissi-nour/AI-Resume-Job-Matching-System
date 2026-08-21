import io
import re

import pdfplumber
from docx import Document


def extract_text_from_pdf(file_bytes: bytes) -> str:
    text_chunks = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text() or ""
            text_chunks.append(page_text)
    return "\n".join(text_chunks)


def extract_text_from_docx(file_bytes: bytes) -> str:
    doc = Document(io.BytesIO(file_bytes))
    paragraphs = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            paragraphs.append(" | ".join(cell.text for cell in row.cells))
    return "\n".join(paragraphs)


def clean_text(text: str) -> str:
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def extract_text(filename: str, file_bytes: bytes) -> str:
    """Dispatch to the right extractor based on file extension."""
    lower = filename.lower()
    if lower.endswith(".pdf"):
        raw = extract_text_from_pdf(file_bytes)
    elif lower.endswith(".docx"):
        raw = extract_text_from_docx(file_bytes)
    elif lower.endswith(".txt"):
        raw = file_bytes.decode("utf-8", errors="ignore")
    else:
        raise ValueError(f"Unsupported file type: {filename}")

    text = clean_text(raw)
    if not text:
        raise ValueError(
            f"No extractable text found in {filename}. "
            "It may be a scanned/image-based document."
        )
    return text
